import os
import re
import shutil
import tempfile
from docx import Document
from lxml import etree

class ElsevierConverter:
    def __init__(self, docx_path, dtd_path, meta_xml_path=None):
        self.docx_path = docx_path
        self.dtd_path = dtd_path
        self.meta_xml_path = meta_xml_path
        
        # Elsevier 5.7.0 Namespaces
        self.NSMAP = {
            None: "http://www.elsevier.com/xml/ja/dtd",
            "ce": "http://www.elsevier.com/xml/common/dtd",
            "sa": "http://www.elsevier.com/xml/common/struct-aff/dtd",
            "sb": "http://www.elsevier.com/xml/common/struct-bib/dtd",
            "xlink": "http://www.w3.org/1999/xlink"
        }

    def convert(self):
        """Reads DOCX and builds a DTD-compliant XML tree."""
        if not os.path.exists(self.docx_path):
            raise FileNotFoundError(f"DOCX not found: {self.docx_path}")

        doc = Document(self.docx_path)
        root = etree.Element("article", nsmap=self.NSMAP)
        root.set("docsubtype", "fla")
        root.set("version", "5.7")

        meta = self._load_meta()

        # --- 1. ITEM-INFO (Mandatory Metadata) ---
        item_info = etree.SubElement(root, "item-info")
        jid = etree.SubElement(item_info, "jid")
        jid.text = meta.get("jid", "CHAOS")
        aid = etree.SubElement(item_info, "aid")
        aid.text = meta.get("aid", "115581")
        if meta.get("article_number"):
            article_number = etree.SubElement(
                item_info, "{http://www.elsevier.com/xml/common/dtd}article-number"
            )
            article_number.text = meta["article_number"]
        pii = etree.SubElement(item_info, "{http://www.elsevier.com/xml/common/dtd}pii")
        pii.text = meta.get("pii", "UNKNOWN")
        if meta.get("doi"):
            doi = etree.SubElement(item_info, "{http://www.elsevier.com/xml/common/dtd}doi")
            doi.text = meta["doi"]

        copyright_el = etree.SubElement(
            item_info, "{http://www.elsevier.com/xml/common/dtd}copyright"
        )
        copyright_el.set("type", meta.get("copyright_type", "unknown"))
        copyright_el.set("year", meta.get("copyright_year", "2024"))
        copyright_el.text = meta.get("copyright_text", "Copyright (c) Elsevier B.V.")
        if meta.get("copyright_line"):
            copyright_line = etree.SubElement(
                item_info, "{http://www.elsevier.com/xml/common/dtd}copyright-line"
            )
            copyright_line.text = meta["copyright_line"]
        
        # --- 2. HEAD (Title/Authors) ---
        head = etree.SubElement(root, "head")
        ce_title = etree.SubElement(head, "{http://www.elsevier.com/xml/common/dtd}title")
        ce_title.text = meta.get("title") or (doc.paragraphs[0].text if doc.paragraphs else "Untitled Document")

        author_group = etree.SubElement(head, "{http://www.elsevier.com/xml/common/dtd}author-group")
        authors = meta.get("authors", [])
        for author in authors:
            author_el = etree.SubElement(author_group, "{http://www.elsevier.com/xml/common/dtd}author")
            if author.get("degrees"):
                degrees = etree.SubElement(author_el, "{http://www.elsevier.com/xml/common/dtd}degrees")
                degrees.text = author["degrees"]
            given = etree.SubElement(author_el, "{http://www.elsevier.com/xml/common/dtd}given-name")
            given.text = author.get("given_name", "")
            surname = etree.SubElement(author_el, "{http://www.elsevier.com/xml/common/dtd}surname")
            surname.text = author.get("surname", "")
            if author.get("email"):
                eaddr = etree.SubElement(author_el, "{http://www.elsevier.com/xml/common/dtd}e-address")
                eaddr.text = author["email"]
        if not authors:
            author_el = etree.SubElement(author_group, "{http://www.elsevier.com/xml/common/dtd}author")
            given = etree.SubElement(author_el, "{http://www.elsevier.com/xml/common/dtd}given-name")
            given.text = "Unknown"
            surname = etree.SubElement(author_el, "{http://www.elsevier.com/xml/common/dtd}surname")
            surname.text = "Author"

        # --- 3. BODY (Content) ---
        body = etree.SubElement(root, "body")
        sections = etree.SubElement(body, "{http://www.elsevier.com/xml/common/dtd}sections")
        current_section = None

        for para in doc.paragraphs[1:]:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name.lower()
            if 'heading' in style:
                current_section = etree.SubElement(
                    sections, "{http://www.elsevier.com/xml/common/dtd}section"
                )
                title = etree.SubElement(current_section, "{http://www.elsevier.com/xml/common/dtd}section-title")
                title.text = text
            else:
                if current_section is None:
                    current_section = etree.SubElement(
                        sections, "{http://www.elsevier.com/xml/common/dtd}section"
                    )
                    title = etree.SubElement(
                        current_section, "{http://www.elsevier.com/xml/common/dtd}section-title"
                    )
                    title.text = "Introduction"
                parent = current_section
                ce_para = etree.SubElement(parent, "{http://www.elsevier.com/xml/common/dtd}para")
                ce_para.text = text
        
        for section in sections.findall("{http://www.elsevier.com/xml/common/dtd}section"):
            if not section.findall("{http://www.elsevier.com/xml/common/dtd}para"):
                filler = etree.SubElement(section, "{http://www.elsevier.com/xml/common/dtd}para")
                filler.text = "Content pending."

        return root

    def _load_meta(self):
        meta = {}
        if not self.meta_xml_path or not os.path.isfile(self.meta_xml_path):
            return meta

        tree = etree.parse(self.meta_xml_path)
        item_info = tree.find(".//item-info")
        if item_info is None:
            return meta

        def _text(tag):
            node = item_info.find(tag)
            return node.text.strip() if node is not None and node.text else ""

        meta["jid"] = _text("jid")
        meta["aid"] = _text("aid")
        meta["pii"] = _text("pii")
        meta["doi"] = _text("doi")
        meta["article_number"] = _text("article-number")
        meta["title"] = _text("item-title")

        accept_date = item_info.find("accept-date/date")
        if accept_date is not None and accept_date.get("yr"):
            meta["copyright_year"] = accept_date.get("yr")

        authors = []
        first_author = item_info.find("first-author")
        if first_author is not None:
            authors.append(
                {
                    "degrees": _text("first-author/degree"),
                    "given_name": _text("first-author/fnm"),
                    "surname": _text("first-author/snm"),
                }
            )

        corr_author = item_info.find("corr-author")
        if corr_author is not None:
            authors.append(
                {
                    "degrees": _text("corr-author/degree"),
                    "given_name": _text("corr-author/fnm"),
                    "surname": _text("corr-author/snm"),
                    "email": _text("corr-author/aff/ead"),
                }
            )

        if authors:
            meta["authors"] = authors

        return meta

    def validate(self, xml_tree):
        """Validates using a robust loader to handle modular DTD files."""
        if not os.path.isfile(self.dtd_path):
            print(f"ERROR: DTD file missing: {self.dtd_path}")
            return False

        # --- Sanity Check: Is it actually HTML? ---
        with open(self.dtd_path, 'r', encoding='utf-8', errors='ignore') as f:
            first_line = f.readline().lower()
            if '<html' in first_line or '<!doctype html' in first_line:
                print("ERROR: Your art570.dtd file appears to be an HTML webpage.")
                print("Please re-download the RAW DTD files from Elsevier.")
                return False

        try:
            print(f"Attempting to parse DTD: {os.path.abspath(self.dtd_path)}")

            # Using external_id=None and file=path is the standard for local files
            # lxml will look for .ent files in the same directory as this path
            dtd = etree.DTD(external_id=None, file=os.path.abspath(self.dtd_path))
            return self._validate_with_dtd(dtd, xml_tree)

        except etree.DTDParseError as e:
            print(f"DTD Parsing Error: {e}")
            print("Retrying with a patched copy of the DTD/ENT files.")
            try:
                return self._validate_with_patched_dtd(xml_tree)
            except etree.DTDParseError as patched_error:
                print(f"DTD Parsing Error (patched copy failed): {patched_error}")
                print("Check common170.ent for parameter entities using dots in their names.")
                return False
        except Exception as e:
            print(f"Error during validation: {e}")
            return False

    def _validate_with_dtd(self, dtd, xml_tree):
        if dtd.validate(xml_tree):
            print("Validation Successful: XML is art570 compliant.")
            return True

        print("Validation Failed. Structural Errors:")
        for error in dtd.error_log.filter_from_errors()[:10]:
            print(f"  - Line {error.line}: {error.message}")
        return False

    def _validate_with_patched_dtd(self, xml_tree):
        source_dir = os.path.dirname(os.path.abspath(self.dtd_path))
        dtd_filename = os.path.basename(self.dtd_path)

        with tempfile.TemporaryDirectory() as temp_dir:
            dtd_source = os.path.join(source_dir, dtd_filename)
            common_source = os.path.join(source_dir, "common170.ent")
            if not os.path.isfile(dtd_source) or not os.path.isfile(common_source):
                raise etree.DTDParseError("Required DTD/ENT files are missing.", 0, 0, 0)

            shutil.copy2(dtd_source, os.path.join(temp_dir, dtd_filename))
            shutil.copy2(common_source, os.path.join(temp_dir, "common170.ent"))

            common_path = os.path.join(temp_dir, "common170.ent")
            with open(common_path, "r", encoding="utf-8") as f:
                content = f.read()
            # Drop MathML DTD inclusion only if the MathML module files are missing.
            qname_mod = os.path.join(source_dir, "mathml3-qname.mod")
            skipped_mathml = False
            if not os.path.isfile(qname_mod):
                content = re.sub(
                    r'<!ENTITY % mathml-dtd[\\s\\S]*?>\\s*%mathml-dtd;\\s*',
                    '<!ENTITY % mathml-dtd \"\" >\\n',
                    content,
                    count=1,
                )
                skipped_mathml = True
            with open(common_path, "w", encoding="utf-8") as f:
                f.write(content)

            if skipped_mathml:
                print("Patched DTD: MathML module skipped for validation.")
            patched_dtd_path = os.path.join(temp_dir, dtd_filename)
            dtd = etree.DTD(external_id=None, file=patched_dtd_path)
            return self._validate_with_dtd(dtd, xml_tree)

    def save_xml(self, xml_tree, output_path):
        doctype = '<!DOCTYPE article PUBLIC "-//ES//DTD journal article DTD version 5.7.0//EN//XML" "art570.dtd">'
        with open(output_path, 'wb') as f:
            f.write(etree.tostring(
                xml_tree, 
                pretty_print=True, 
                xml_declaration=True, 
                encoding="UTF-8", 
                doctype=doctype
            ))
        print(f"XML saved to: {output_path}")

if __name__ == "__main__":
    # Settings
    INPUT_DOCX = "CHAOS_115581_original.docx"
    OUTPUT_XML = "output_article.xml"
    DTD_FILE = "art570.dtd" 
    META_XML = "CHAOS_115581_meta_final.xml"

    converter = ElsevierConverter(INPUT_DOCX, DTD_FILE, META_XML)

    try:
        # Step 1: Conversion
        xml_result = converter.convert()

        # Step 2: Save to Disk
        converter.save_xml(xml_result, OUTPUT_XML)

        # Step 3: DTD Comparison
        converter.validate(xml_result)

    except Exception as e:
        print(f"Process stopped: {e}")
